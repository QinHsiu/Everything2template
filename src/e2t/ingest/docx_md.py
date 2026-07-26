"""DOCX / Markdown / plain text → CIR."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from e2t.cir import SourceKind
from e2t.ingest.draft import draft_cir_from_text
from e2t.paths import display_path


def ingest_text_file(path: str | Path):
    path = Path(path)
    text = path.read_text(encoding="utf-8", errors="replace")
    return draft_cir_from_text(
        text,
        source_kind=SourceKind.text,
        source_ref=display_path(path),
        title=None,
    )


def ingest_markdown(path: str | Path):
    path = Path(path)
    text = path.read_text(encoding="utf-8", errors="replace")
    return draft_cir_from_text(
        text,
        source_kind=SourceKind.markdown,
        source_ref=display_path(path),
        title=None,
    )


def ingest_docx(path: str | Path):
    path = Path(path)
    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n\n".join(parts)
    return draft_cir_from_text(
        text or "(empty docx)",
        source_kind=SourceKind.docx,
        source_ref=display_path(path),
        title=path.stem,
    )
